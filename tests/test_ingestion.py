import pathlib

import docx
import pytest

import careshield.contracts.schemas as schemas
import careshield.retrieval.ingestion as ingestion


def test_parse_text_document_bytes() -> None:
    """Verify text parsing."""
    parsed = ingestion.parse_document_bytes(
        content=b"Clinical report says vendor sharing requires redaction and approval.",
        source_name="report.txt",
    )
    assert parsed.parser == "utf8-text"
    assert "vendor sharing" in parsed.text


def test_parse_docx_document_file(tmp_path: pathlib.Path) -> None:
    """Verify Word parsing."""
    path = tmp_path / "care-report.docx"
    document = docx.Document()
    document.add_paragraph("Clinical report section.")
    document.add_paragraph("Vendor sharing requires de-identification and audit trail.")
    document.save(str(path))

    parsed = ingestion.parse_document_file(path=path)

    assert parsed.parser == "python-docx"
    assert "Vendor sharing requires de-identification" in parsed.text


def test_parse_pdf_document_bytes() -> None:
    """Verify PDF parsing."""
    parsed = ingestion.parse_document_bytes(
        content=_minimal_pdf_bytes(text="Vendor sharing requires redaction and compliance approval."),
        source_name="care-report.pdf",
    )

    assert parsed.parser == "pypdf"
    assert "Vendor sharing requires redaction" in parsed.text


def test_unsupported_document_type_is_rejected() -> None:
    """Verify unsupported parser types fail clearly."""
    with pytest.raises(expected_exception=ingestion.DocumentParseError):
        ingestion.parse_document_bytes(content=b"hello world", source_name="report.csv")


def test_chunking_and_document_building_preserve_policy_metadata() -> None:
    """Verify chunks keep sensitivity and role metadata."""
    chunks = ingestion.chunk_text(
        text=" ".join(["redaction"] * 220),
        max_words=50,
        overlap_words=10,
    )
    documents = ingestion.build_documents_from_text(
        text="Clinical report requires redaction before external sharing. " * 20,
        source_name="care-report.md",
        sensitivity=schemas.Sensitivity.clinical,
        max_words=30,
        overlap_words=5,
    )

    assert len(chunks) > 1
    assert len(documents) > 1
    assert documents[0].sensitivity == schemas.Sensitivity.clinical
    assert schemas.Role.nurse in documents[0].allowed_roles
    assert schemas.Role.external_vendor not in documents[0].allowed_roles


def _minimal_pdf_bytes(*, text: str) -> bytes:
    """Build a tiny parseable PDF for offline tests.

    :param text: Text to place on the PDF page.
    :return: PDF bytes.
    """
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("utf-8")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
            b"/MediaBox [0 0 612 792] /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("utf-8") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("utf-8"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("utf-8"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("utf-8"))
    pdf.extend(
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\nstartxref\n{xref_offset}\n%%EOF\n".encode(
            "utf-8"
        )
    )
    return bytes(pdf)
