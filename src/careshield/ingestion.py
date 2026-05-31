from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from careshield.schemas import Document, Role, Sensitivity


class DocumentParseError(ValueError):
    """Raised when an uploaded knowledge document cannot be parsed safely."""


@dataclass(frozen=True)
class ParsedDocument:
    source_name: str
    parser: str
    text: str


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def parse_document_bytes(content: bytes, source_name: str) -> ParsedDocument:
    suffix = Path(source_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentParseError(f"unsupported document type {suffix!r}; supported: {supported}")
    if not content:
        raise DocumentParseError("document is empty")

    if suffix in {".txt", ".md"}:
        text = content.decode("utf-8", errors="replace")
        parser = "utf8-text"
    elif suffix == ".docx":
        text = _parse_docx(content)
        parser = "python-docx"
    else:
        text = _parse_pdf(content)
        parser = "pypdf"

    normalized = _normalize_text(text)
    if len(normalized) < 20:
        raise DocumentParseError("document did not contain enough readable text")
    return ParsedDocument(source_name=source_name, parser=parser, text=normalized)


def parse_document_file(path: str | Path) -> ParsedDocument:
    document_path = Path(path)
    return parse_document_bytes(document_path.read_bytes(), document_path.name)


def chunk_text(text: str, max_words: int = 90, overlap_words: int = 18) -> list[str]:
    if max_words <= 0:
        raise ValueError("max_words must be positive")
    if overlap_words < 0 or overlap_words >= max_words:
        raise ValueError("overlap_words must be smaller than max_words")

    words = text.split()
    if not words:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + max_words, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = end - overlap_words
    return chunks


def build_documents_from_text(
    text: str,
    source_name: str,
    sensitivity: Sensitivity,
    max_words: int = 90,
    overlap_words: int = 18,
) -> list[Document]:
    source_id = _safe_id(Path(source_name).stem)
    allowed_roles = _allowed_roles_for_sensitivity(sensitivity)
    chunks = chunk_text(text, max_words=max_words, overlap_words=overlap_words)
    return [
        Document(
            id=f"{source_id}-chunk-{index}",
            title=f"{Path(source_name).stem} section {index}",
            body=chunk,
            sensitivity=sensitivity,
            allowed_roles=allowed_roles,
            tags=["uploaded-report", Path(source_name).suffix.lower().lstrip(".") or "text"],
        )
        for index, chunk in enumerate(chunks, start=1)
    ]


def _parse_docx(content: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:  # pragma: no cover - library exception types vary
        raise DocumentParseError("could not read docx document") from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pragma: no cover - library exception types vary
        raise DocumentParseError("could not read pdf document") from exc
    return "\n".join(pages)


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _safe_id(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return normalized or "document"


def _allowed_roles_for_sensitivity(sensitivity: Sensitivity) -> list[Role]:
    if sensitivity == Sensitivity.public:
        return list(Role)
    if sensitivity == Sensitivity.internal:
        return [
            Role.doctor,
            Role.nurse,
            Role.billing_analyst,
            Role.vendor_manager,
            Role.compliance_officer,
        ]
    if sensitivity == Sensitivity.clinical:
        return [Role.doctor, Role.nurse, Role.compliance_officer]
    if sensitivity == Sensitivity.billing:
        return [Role.billing_analyst, Role.compliance_officer]
    return [Role.compliance_officer]
