import io
import pathlib
import re

import docx
import pydantic
import pypdf

from careshield import contracts


class DocumentParseError(ValueError):
    """Raised when an uploaded knowledge document cannot be parsed safely."""


class ParsedDocument(pydantic.BaseModel):
    """Parsed text and parser metadata for an uploaded file."""

    model_config = pydantic.ConfigDict(
        frozen=True,
        json_schema_extra={
            "examples": [
                {
                    "source_name": "synthetic-care-report.md",
                    "parser": "utf8-text",
                    "text": "Vendor sharing requires de-identification and approval.",
                }
            ]
        },
    )

    source_name: str = pydantic.Field(min_length=1)
    parser: str = pydantic.Field(min_length=1)
    text: str = pydantic.Field(min_length=20)


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def parse_document_bytes(*, content: bytes, source_name: str) -> ParsedDocument:
    """Parse supported document bytes into normalized text.

    :param content: Raw file bytes from disk or upload.
    :param source_name: Original file name used to infer parser type.
    :return: Parsed text with parser metadata.
    """
    suffix = pathlib.Path(source_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentParseError(f"unsupported document type {suffix!r}; supported: {supported}")
    if not content:
        raise DocumentParseError("document is empty")

    if suffix in {".txt", ".md"}:
        text = content.decode(encoding="utf-8", errors="replace")
        parser = "utf8-text"
    elif suffix == ".docx":
        text = _parse_docx(content=content)
        parser = "python-docx"
    else:
        text = _parse_pdf(content=content)
        parser = "pypdf"

    # Normalize parser output before chunking so retrieval behaves similarly
    # across Markdown, PDF, and Word inputs.
    normalized = _normalize_text(text=text)
    if len(normalized) < 20:
        raise DocumentParseError("document did not contain enough readable text")
    return ParsedDocument(source_name=source_name, parser=parser, text=normalized)


def parse_document_file(*, path: str | pathlib.Path) -> ParsedDocument:
    """Parse a document from the local filesystem.

    :param path: Local file path.
    :return: Parsed text with parser metadata.
    """
    document_path = pathlib.Path(path)
    return parse_document_bytes(content=document_path.read_bytes(), source_name=document_path.name)


def chunk_text(*, text: str, max_words: int = 90, overlap_words: int = 18) -> list[str]:
    """Split text into overlapping word chunks.

    :param text: Text to chunk.
    :param max_words: Maximum words in each chunk.
    :param overlap_words: Words repeated between neighboring chunks.
    :return: Ordered text chunks.
    """
    if max_words <= 0:
        raise ValueError("max_words must be positive")
    if overlap_words < 0 or overlap_words >= max_words:
        raise ValueError("overlap_words must be smaller than max_words")

    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + max_words, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = end - overlap_words
    return chunks


def build_documents_from_text(
    *,
    text: str,
    source_name: str,
    sensitivity: contracts.schema.Sensitivity,
    max_words: int = 90,
    overlap_words: int = 18,
) -> list[contracts.schema.Document]:
    """Turn parsed text into policy-aware retrievable chunks.

    :param text: Parsed document text.
    :param source_name: Source file name.
    :param sensitivity: Sensitivity assigned to every chunk.
    :param max_words: Maximum words in each chunk.
    :param overlap_words: Words repeated between neighboring chunks.
    :return: Documents that can be indexed and retrieved.
    """
    source_path = pathlib.Path(source_name)
    source_id = _safe_id(value=source_path.stem)
    allowed_roles = _allowed_roles_for_sensitivity(sensitivity=sensitivity)
    chunks = chunk_text(text=text, max_words=max_words, overlap_words=overlap_words)

    # Every chunk carries policy metadata because retrieval must be able to
    # filter evidence before any model prompt is built.
    return [
        contracts.schema.Document(
            id=f"{source_id}-chunk-{index}",
            title=f"{source_path.stem} section {index}",
            body=chunk,
            sensitivity=sensitivity,
            allowed_roles=allowed_roles,
            tags=["uploaded-report", source_path.suffix.lower().lstrip(".") or "text"],
        )
        for index, chunk in enumerate(chunks, start=1)
    ]


def _parse_docx(*, content: bytes) -> str:
    """Extract text from a Word document.

    :param content: DOCX bytes.
    :return: Extracted paragraph and table text.
    """
    try:
        document = docx.Document(docx=io.BytesIO(content))
    except Exception as exc:  # pragma: no cover - library exception types vary
        raise DocumentParseError("could not read docx document") from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(*, content: bytes) -> str:
    """Extract text from a PDF document.

    :param content: PDF bytes.
    :return: Extracted page text.
    """
    try:
        reader = pypdf.PdfReader(stream=io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pragma: no cover - library exception types vary
        raise DocumentParseError("could not read pdf document") from exc
    return "\n".join(pages)


def _normalize_text(*, text: str) -> str:
    """Normalize whitespace and null characters.

    :param text: Raw extracted text.
    :return: Clean text ready for chunking.
    """
    text = text.replace("\x00", " ")
    text = re.sub(pattern=r"\s+", repl=" ", string=text)
    return text.strip()


def _safe_id(*, value: str) -> str:
    """Build a stable ID fragment from a source name.

    :param value: Raw file stem.
    :return: URL-safe lowercase identifier.
    """
    normalized = re.sub(pattern=r"[^a-z0-9]+", repl="-", string=value.lower()).strip("-")
    return normalized or "document"


def _allowed_roles_for_sensitivity(
    *,
    sensitivity: contracts.schema.Sensitivity,
) -> list[contracts.schema.Role]:
    """Derive allowed roles for uploaded chunks from sensitivity.

    :param sensitivity: Sensitivity assigned by the upload request.
    :return: Roles allowed to retrieve the uploaded chunks.
    """
    if sensitivity == contracts.schema.Sensitivity.public:
        return list(contracts.schema.Role)
    if sensitivity == contracts.schema.Sensitivity.internal:
        return [
            contracts.schema.Role.doctor,
            contracts.schema.Role.nurse,
            contracts.schema.Role.billing_analyst,
            contracts.schema.Role.vendor_manager,
            contracts.schema.Role.compliance_officer,
        ]
    if sensitivity == contracts.schema.Sensitivity.clinical:
        return [
            contracts.schema.Role.doctor,
            contracts.schema.Role.nurse,
            contracts.schema.Role.compliance_officer,
        ]
    if sensitivity == contracts.schema.Sensitivity.billing:
        return [contracts.schema.Role.billing_analyst, contracts.schema.Role.compliance_officer]
    return [contracts.schema.Role.compliance_officer]
