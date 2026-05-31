import pathlib

import docx

ROOT = pathlib.Path(__file__).resolve().parents[1]
EXAMPLES_DIR = ROOT / "examples"
REPORT_STEM = "synthetic-care-report"
REPORT_TEXT = """
Synthetic Care Sharing Report

Care coordination teams may prepare a de-identified discharge summary for an
approved external vendor only after compliance approval.

The report contains synthetic sensitive fields: Patient Jane Example,
jane.example@example.invalid, +1-555-0100, MRN MRN-000-EXAMPLE, insurance ID
INS-000-EXAMPLE, and diagnosis asthma.

Before vendor sharing, teams must redact patient names, phone numbers, email
addresses, medical record numbers, insurance identifiers, and diagnosis details.

The approved model gateway must validate structured responses, keep audit
traces, cite retrieved evidence, and avoid sending restricted fields to public
models.

Compliance officers may review full synthetic reports. Nurses and doctors may
use clinical guidance. External vendors may receive only approved public or
de-identified summaries.
""".strip()


def main() -> None:
    """Generate synthetic Markdown, Word, and PDF examples."""
    EXAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    _write_markdown(path=EXAMPLES_DIR / f"{REPORT_STEM}.md")
    _write_docx(path=EXAMPLES_DIR / f"{REPORT_STEM}.docx")
    _write_pdf(path=EXAMPLES_DIR / f"{REPORT_STEM}.pdf")


def _write_markdown(*, path: pathlib.Path) -> None:
    """Write the Markdown example.

    :param path: Output path.
    """
    path.write_text(data=f"# {REPORT_TEXT}\n", encoding="utf-8")


def _write_docx(*, path: pathlib.Path) -> None:
    """Write the Word example.

    :param path: Output path.
    """
    document = docx.Document()
    for index, paragraph in enumerate(REPORT_TEXT.split("\n\n")):
        if index == 0:
            document.add_heading(text=paragraph, level=1)
            continue
        document.add_paragraph(text=paragraph)
    document.save(str(path))


def _write_pdf(*, path: pathlib.Path) -> None:
    """Write a small parseable PDF example.

    :param path: Output path.
    """
    path.write_bytes(data=_minimal_pdf_bytes(text=REPORT_TEXT.replace("\n", " ")))


def _minimal_pdf_bytes(*, text: str) -> bytes:
    """Build a tiny text-only PDF without adding a PDF dependency.

    :param text: Text to place on the PDF page.
    :return: PDF bytes.
    """
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    wrapped_lines = _wrap_text(text=escaped_text, width=82)
    commands = ["BT /F1 10 Tf 72 740 Td 14 TL"]
    commands.extend(f"({line}) Tj T*" for line in wrapped_lines[:42])
    commands.append("ET")
    stream = "\n".join(commands).encode(encoding="utf-8")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
            b"/MediaBox [0 0 612 792] /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("utf-8") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode())
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    pdf.extend(
        (f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\nstartxref\n{xref_offset}\n%%EOF\n").encode()
    )
    return bytes(pdf)


def _wrap_text(*, text: str, width: int) -> list[str]:
    """Wrap text into PDF-friendly lines.

    :param text: Text to wrap.
    :param width: Maximum characters per line.
    :return: Wrapped lines.
    """
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join([*current, word])
        if len(candidate) > width and current:
            lines.append(" ".join(current))
            current = [word]
            continue
        current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


if __name__ == "__main__":
    main()
